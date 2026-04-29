"""DOCX (.docx) → markdown.

Опционально: использует python-docx если установлен.
"""
from pathlib import Path

try:
    from docx import Document as DocxDocument
except ImportError:
    raise ImportError("Для ингестии DOCX установите: pip install python-docx")

from docstoolkit.ingest.base import Document, IngestError, Source


def ingest(path: Path) -> Document:
    try:
        docx_doc = DocxDocument(str(path))
    except Exception as e:
        raise IngestError(f"DOCX parse error: {e}")

    parts: list[str] = []
    title = ""

    # Core properties
    cp = docx_doc.core_properties
    if cp.title:
        title = cp.title

    for para in docx_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style else "").lower()
        if "heading 1" in style or style == "title":
            if not title:
                title = text
            parts.append(f"# {text}")
        elif "heading 2" in style:
            parts.append(f"## {text}")
        elif "heading 3" in style:
            parts.append(f"### {text}")
        elif "heading 4" in style:
            parts.append(f"#### {text}")
        elif "heading 5" in style:
            parts.append(f"##### {text}")
        elif "heading 6" in style:
            parts.append(f"###### {text}")
        elif "list" in style or "bullet" in style:
            parts.append(f"- {text}")
        else:
            parts.append(text)

    # Таблицы
    for table in docx_doc.tables:
        rows = []
        for r, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if r == 0:
                rows.append("|" + "|".join("---" for _ in cells) + "|")
        parts.append("\n".join(rows))

    if not parts:
        raise IngestError("DOCX: не извлечено контента")

    content = "\n\n".join(parts)

    return Document(
        title=title or path.stem,
        content=content,
        source=Source.from_path(path, "docx"),
        metadata={
            "author": cp.author or "",
            "paragraph_count": len(docx_doc.paragraphs),
            "table_count": len(docx_doc.tables),
        },
    )
